import os 
import json 
import numpy as np 
import pandas as pd
import jieba 
import re
from typing import List, Dict, Any
from docx import Document
import fitz
from bs4 import BeautifulSoup
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import matplotlib.pyplot as plt
import seaborn as sns
from wordcloud import WordCloud
from collections import Counter
import matplotlib
from datetime import datetime
import networkx as nx
from itertools import combinations
from scipy.spatial.distance import cosine
from sklearn.decomposition import TruncatedSVD
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import Normalizer


matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Arial Unicode MS', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False
from sklearn.metrics.pairwise import cosine_similarity

stopwords_path  = r'server/src/data/stopwords.txt'

## 读取传回来的json文件中的路径，并且将路径中的文件读出文本，
## 
class Get_file:
    def __init__(self, path: str, is_json: bool = True):
        """
        初始化文件处理器
        
        Args:
            path: 路径 (可以是文件夹路径或JSON配置文件路径)
            is_json: 是否为JSON配置文件
        """
        if is_json:
            # JSON配置文件模式
            self.json_path = path
            with open(self.json_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
            self.json_dir = config.get('json_path', '')
        else:
            # 直接文件夹路径模式
            self.json_dir = path
        
        self.file_path = []  # 文件路径
        self.file_name = []  # 文件名
        self.kind = ['']
        self.all_content = []  # 全文件文本
        self.processed_files_content: Dict[str, str] = {}  # 文件名对应的文件内容
        self.word_dict = {}  # 词频统计
        self.inverted_index = {}  # 倒排索引
        self.file_word_count = {}  # 文件词频统计
        self.content_list = []
        self.total_docs = 0
        self.stopwords = self._load_stopwords()
    # 读取文本
    def read_content(self, kind: str,file_path: str) -> str:
        """读取纯文本文件 (.txt) 并返回其内容。"""
        if kind == 'txt':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return content
            except FileNotFoundError:
                return f"错误: 文件 '{file_path}' 未找到。"
            except Exception as e:
                return f"读取TXT文件时发生错误: {e}"
        elif kind == 'json':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                # 将解析后的Python对象转换回格式化的JSON字符串
                content = json.dumps(data, ensure_ascii=False, indent=4)
                return content
            except FileNotFoundError:
                return f"错误: 文件 '{file_path}' 未找到。"
            except json.JSONDecodeError:
                return f"错误: 文件 '{file_path}' 不是有效的JSON格式。"
            except Exception as e:
                return f"读取JSON文件时发生错误: {e}"
        elif kind == 'csv':
            try:
                # encoding参数对于read_csv是有效的
                df = pd.read_csv(file_path, encoding='utf-8')
                content = df.to_string()
                return content
            except FileNotFoundError:
                return f"错误: 文件 '{file_path}' 未找到。"
            except pd.errors.EmptyDataError:
                return f"错误: CSV文件 '{file_path}' 为空。"
            except Exception as e:
                return f"读取CSV文件时发生错误: {e}"
        elif kind == 'xlsx' or kind == 'xls':
            try:
                excel_data = pd.read_excel(file_path, sheet_name=None)
                all_sheets_content = []
                if isinstance(excel_data, dict): # 如果有多个工作表
                    for sheet_name, df in excel_data.items():
                        sheet_content = f"--- 内容来自工作表: {sheet_name} ---\n"
                        sheet_content += df.to_string()
                        all_sheets_content.append(sheet_content)
                else: # 如果只有一个工作表，excel_data直接是DataFrame
                    all_sheets_content.append(excel_data.to_string())
                return "\n\n".join(all_sheets_content)
            except FileNotFoundError:
                return f"错误: 文件 '{file_path}' 未找到。"
            except Exception as e:
                return f"读取Excel文件时发生错误: {e}"
        elif kind == 'html':
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                # 使用BeautifulSoup解析HTML并提取文本
                soup = BeautifulSoup(html_content, 'html.parser')
                # 移除script和style标签
                for script in soup(["script", "style"]):
                    script.decompose()
                # 提取纯文本
                content = soup.get_text()
                # 清理多余的空白字符
                content = re.sub(r'\n\s*\n', '\n\n', content)
                content = re.sub(r' +', ' ', content)
                return content.strip()
            except FileNotFoundError:
                return f"错误: 文件 '{file_path}' 未找到。"
            except Exception as e:
                return f"读取HTML文件时发生错误: {e}"
        elif kind == 'docx':
            try:
                doc = Document(file_path)
                content_parts = []
                # 读取所有段落
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content_parts.append(paragraph.text.strip())

                # 读取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            content_parts.append(' | '.join(row_text))

                content = '\n'.join(content_parts)
                return content if content else "文档内容为空"
            except FileNotFoundError:
                return f"错误: 文件 '{file_path}' 未找到。"
            except Exception as e:
                return f"读取DOCX文件时发生错误: {e}"
        elif kind == 'pdf':
            try:
                # 使用PyMuPDF (fitz) 读取PDF文件
                pdf_document = fitz.open(file_path)
                content_parts = []

                # 遍历所有页面
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                    # 提取页面文本
                    page_text = page.get_text()
                    if page_text.strip():
                        content_parts.append(f"--- 第{page_num + 1}页 ---")
                        content_parts.append(page_text.strip())

                pdf_document.close()

                content = '\n\n'.join(content_parts)
                return content if content else "PDF文档内容为空或无法提取文本"

            except FileNotFoundError:
                return f"错误: 文件 '{file_path}' 未找到。"
            except Exception as e:
                return f"读取PDF文件时发生错误: {e}"

        else :
            return f"错误: 未知或不支持的文件类型 '{kind}'"  # 返回错误信息而不是空值
    #洁词
    def process_text(self, content: list) -> List[str]:
        process_content = []
        for text in content:
            text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s]', '', text)
            words = list(jieba.cut(text))
            filtered_words = [
                word.strip() for word in words 
                if word.strip() and word.strip() not in self.stopwords and len(word.strip()) > 1
            ]
            processed_text = ' '.join(filtered_words)
            process_content.append(processed_text)

        return process_content


    def _load_stopwords(self) -> List[str]:
        """加载停用词"""
        try:
            with open(stopwords_path, 'r', encoding='utf-8') as f:
                print(f"加载停用词文件: {stopwords_path}")
                return [word.strip() for word in f.readlines()]
            
        except FileNotFoundError:
            return ['的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个']

    #统计词频
    def all_word_count(self):
        word_dict = {}
        for file_name, content in self.processed_files_content.items():
            if content.startswith("错误:"):
                continue
            process_content = self.process_text([content])
            if process_content:
                word_list = process_content[0].split()
                for word in word_list:
                    word_dict[word] = word_dict.get(word, 0) + 1
        sorted_word_dict = sorted(word_dict.items(), key=lambda x: x[1], reverse=True)
        self.word_dict = {word: count for word, count in sorted_word_dict[:100]}  # 只保留前100个高频词
        return self.word_dict

    #构建倒索排序
    def build_inverted_index(self) -> None:
        self.inverted_index = {}
        self.file_word_count = {}
        self.total_docs = 0
        for file_name, content in self.processed_files_content.items():
            if content.startswith("错误:"):
                continue
            processed_content = self.process_text([content])
            if not processed_content:
                continue
            words = processed_content[0].split()
            self.total_docs += 1  # 记录文档总数

            doc_word_count = {}
            for pos, word in enumerate(words):
                if not word:
                    continue

                doc_word_count[word] = doc_word_count.get(word, 0) + 1
          
                if word not in self.inverted_index:
                    self.inverted_index[word] = {
                        'doc_freq': 0,  
                        'postings': {}  
                    }
          
                if file_name not in self.inverted_index[word]['postings']:
                    self.inverted_index[word]['postings'][file_name] = {
                        'freq': 0,
                        'positions': []
                    }
                    self.inverted_index[word]['doc_freq'] += 1 
                posting = self.inverted_index[word]['postings'][file_name]
                posting['freq'] += 1
                posting['positions'].append(pos) 
            
            self.file_word_count[file_name] = doc_word_count

        #读文件   
    def read_dir_file(self):
        for root, dirs, files in os.walk(self.json_dir):
            for file_name in files: # 使用 file_name 避免与外部可能的 file 变量冲突
                self.file_name.append(file_name)
                self.file_path.append(os.path.join(root,file_name))
        for file_path_item in self.file_path: # 遍历收集到的文件路径
                content_or_error_msg = f"错误: 未知或不支持的文件类型 '{os.path.basename(file_path_item)}'." # 默认消息
                file_name_extension = os.path.splitext(file_path_item)[1][1:]
                content_or_error_msg = self.read_content(file_name_extension, file_path_item)
                self.processed_files_content[file_path_item] = content_or_error_msg
                self.content_list.append(content_or_error_msg)
    def init_read_dir(self):
        self.read_dir_file()
        self.all_word_count()
        self.build_inverted_index()


class scearch_content:
    def __init__(self):
        """初始化搜索类"""
        self.stopwords = self._load_stopwords()   
    def _load_stopwords(self) -> List[str]:
        """加载停用词"""
        try:
            with open(stopwords_path, 'r', encoding='utf-8') as f:
                return [word.strip() for word in f.readlines()]
        except FileNotFoundError:
            return ['的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个']
    def preprocess_query(self, query: str) -> List[str]:
        """预处理查询文本"""
        # 去除特殊字符，保留中文、英文、数字
        query = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s]', '', query)
        # 分词
        words = list(jieba.cut(query))
        # 去停用词和空词
        words = [word.strip() for word in words 
                if word.strip() and word.strip() not in self.stopwords and len(word.strip()) > 1]
        return words
    def generate_content_preview(self, content: str, keyword: str, preview_length: int = 200) -> str:
        """生成包含关键词的内容预览"""
        if not content or not keyword:
            return "无预览内容"

        # 查找关键词位置
        keyword_pos = content.find(keyword)
        if keyword_pos == -1:
            # 如果没找到关键词，返回开头部分
            return content[:preview_length] + "..." if len(content) > preview_length else content

        # 计算预览开始位置
        start_pos = max(0, keyword_pos - preview_length // 2)
        end_pos = min(len(content), start_pos + preview_length)

        preview = content[start_pos:end_pos]

        # 添加省略号
        if start_pos > 0:
            preview = "..." + preview
        if end_pos < len(content):
            preview = preview + "..."
    
        return preview
    def _calculate_textrank_scores(self, words: List[str], window_size: int = 3, damping: float = 0.85, max_iter: int = 100, tolerance: float = 1e-4) -> Dict[str, float]:
        """
        计算词汇的TextRank分数

        Args:
            words: 预处理后的词汇列表
            window_size: 共现窗口大小
            damping: 阻尼系数
            max_iter: 最大迭代次数
            tolerance: 收敛阈值

        Returns:
            词汇TextRank分数字典
        """
        if len(words) < 2:
            return {word: 1.0 for word in words}

        # 构建词汇共现图
        word_graph = self._build_word_graph(words, window_size)

        if not word_graph.nodes():
            return {word: 1.0 for word in set(words)}

        # 初始化分数
        unique_words = list(word_graph.nodes())
        scores = {word: 1.0 for word in unique_words}

        # 迭代计算TextRank分数
        for iteration in range(max_iter):
            new_scores = {}

            for word in unique_words:
                # 计算入链权重和
                in_weight_sum = 0
                neighbors = list(word_graph.neighbors(word))

                for neighbor in neighbors:
                    neighbor_out_degree = word_graph.degree(neighbor)
                    if neighbor_out_degree > 0:
                        edge_weight = word_graph[neighbor][word].get('weight', 1.0)
                        in_weight_sum += (scores[neighbor] * edge_weight) / neighbor_out_degree

                # TextRank公式
                new_scores[word] = (1 - damping) + damping * in_weight_sum

            # 检查收敛
            converged = True
            for word in unique_words:
                if abs(new_scores[word] - scores[word]) > tolerance:
                    converged = False
                    break
                
            scores = new_scores

            if converged:
                break
            
        # 归一化分数
        max_score = max(scores.values()) if scores else 1.0
        if max_score > 0:
            scores = {word: score / max_score for word, score in scores.items()}

        return scores

    def _build_word_graph(self, words: List[str], window_size: int) -> nx.Graph:
        """
        构建词汇共现图

        Args:
            words: 词汇列表
            window_size: 共现窗口大小

        Returns:
            词汇共现图
        """
        graph = nx.Graph()

        # 添加节点
        unique_words = list(set(words))
        graph.add_nodes_from(unique_words)

        # 计算词汇共现关系
        word_pairs = {}

        for i in range(len(words)):
            for j in range(i + 1, min(i + window_size + 1, len(words))):
                word1, word2 = words[i], words[j]
                if word1 != word2:
                    pair = tuple(sorted([word1, word2]))
                    word_pairs[pair] = word_pairs.get(pair, 0) + 1

        # 添加边和权重
        for (word1, word2), weight in word_pairs.items():
            if weight > 0:
                graph.add_edge(word1, word2, weight=weight)

        return graph

    def _calculate_sentence_similarity(self, sent1: List[str], sent2: List[str]) -> float:
        """
        计算句子相似度（用于文档级TextRank）

        Args:
            sent1: 第一个句子的词汇列表
            sent2: 第二个句子的词汇列表

        Returns:
            相似度分数
        """
        if not sent1 or not sent2:
            return 0.0

        # 计算词汇交集
        common_words = set(sent1) & set(sent2)

        if not common_words:
            return 0.0

        # 使用Jaccard相似度
        union_words = set(sent1) | set(sent2)
        similarity = len(common_words) / len(union_words)

        return similarity
    def _get_top_lsi_terms(self, lsi_pipeline, feature_names, doc_vector, top_n: int = 5) -> List[Dict[str, Any]]:
        """
        获取文档在LSI空间中的主要语义成分

        Args:
            lsi_pipeline: 训练好的LSI pipeline
            feature_names: 特征名称列表
            doc_vector: 文档在LSI空间的向量表示
            top_n: 返回前N个主要成分

        Returns:
            主要LSI成分列表
        """
        try:
            # 获取SVD组件
            svd = lsi_pipeline.named_steps['svd']

            # 获取最重要的LSI维度
            top_dimensions = abs(doc_vector).argsort()[::-1][:top_n]

            lsi_components = []
            for dim_idx in top_dimensions:
                if dim_idx < len(svd.components_):
                    # 获取该维度对应的词汇权重
                    component = svd.components_[dim_idx]
                    top_terms_idx = abs(component).argsort()[::-1][:5]

                    terms_info = []
                    for term_idx in top_terms_idx:
                        if term_idx < len(feature_names):
                            terms_info.append({
                                'term': feature_names[term_idx],
                                'weight': round(component[term_idx], 4)
                            })

                    lsi_components.append({
                        'dimension': int(dim_idx),
                        'doc_weight': round(doc_vector[dim_idx], 4),
                        'top_terms': terms_info
                    })

            return lsi_components
        except Exception as e:
                    print(f"获取LSI成分时发生错误: {e}")
                    return []
    def _preprocess_for_lsi(self, content: str) -> str:
            """
            专门为LSI算法预处理文本
            Args:
                content: 原始文本内容
            Returns:
                预处理后的文本
            """
            # 去除特殊字符，保留中文、英文、数字
            content = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s]', ' ', content)

            # 分词
            words = list(jieba.cut(content))

            # 去停用词和短词，保留有意义的词汇
            filtered_words = [
                word.strip() for word in words 
                if (word.strip() and 
                    word.strip() not in self.stopwords and 
                    len(word.strip()) >= 2 and
                    not word.strip().isdigit())  # 过滤纯数字
            ]

            return ' '.join(filtered_words)


    # 词频检索
    def textrank_search(self, query: str, inverted_index: Dict[str, Any], processed_files_content: Dict[str, str], file_word_count: Dict[str, Any], window_size: int = 3, damping: float = 0.85) -> Dict[str, Any]:
        """
        基于TextRank算法的文档检索
        Args:
            query: 搜索查询词
            inverted_index: 倒排索引
            processed_files_content: 文件内容字典
            file_word_count: 文件词频统计
            window_size: 共现窗口大小
            damping: 阻尼系数
        Returns:
            检索结果字典
        """
        if not query:
            return {'found': False, 'message': '查询词为空', 'results': []}

        query_words = self.preprocess_query(query)
        if not query_words:
            return {'found': False, 'message': '查询词无效', 'results': []}

        results = []

        # 为每个文档计算TextRank分数
        for file_name, content in processed_files_content.items():
            if content.startswith("错误:"):
                continue

            # 预处理文档内容
            processed_content = self.preprocess_query(content)
            if not processed_content:
                continue

            # 检查文档是否包含查询词
            doc_query_words = [word for word in query_words if word in processed_content]
            if not doc_query_words:
                continue

            # 计算TextRank分数
            textrank_scores = self._calculate_textrank_scores(processed_content, window_size, damping)

            # 计算查询词的综合分数
            query_score = 0
            matched_words = []

            for word in doc_query_words:
                if word in textrank_scores:
                    word_score = textrank_scores[word]
                    query_score += word_score
                    matched_words.append({
                        'word': word,
                        'textrank_score': round(word_score, 4),
                        'frequency': processed_content.count(word)
                    })

            if query_score > 0:
                # 生成内容预览
                preview = self.generate_content_preview(content, doc_query_words[0])

                # 计算文档相关性分数
                doc_relevance = query_score / len(doc_query_words) if doc_query_words else 0

                results.append({
                    'file_name': file_name,
                    'textrank_score': round(query_score, 4),
                    'relevance_score': round(doc_relevance, 4),
                    'matched_words': matched_words,
                    'content_preview': preview,
                    'total_matched_terms': len(matched_words),
                    'algorithm': 'TextRank'
                })

        # 按TextRank分数排序
        results.sort(key=lambda x: x['textrank_score'], reverse=True)

        return {
            'found': len(results) > 0,
            'query': query,
            'total_results': len(results),
            'results': results,
            'message': f'TextRank检索找到 {len(results)} 个匹配文档' if results else 'TextRank检索未找到匹配文档'
        }

    # tfidf词向量检索，返回对应词每个文章中的查询词向量，以及文件路径 
    def tfidf_search(self, query: str, inverted_index: Dict[str, Any], processed_files_content: Dict[str, str], file_word_count: Dict[str, any], all_content: list) -> Dict[str, any]:
        """
        基于TF-IDF的检索
        Args:
            query: 搜索查询词
            inverted_index: 倒排索引
            processed_files_content: 文件内容字典
            file_word_count: 文件词频统计
            all_content: 所有文件内容列表
        Returns:
            检索结果字典
        """
        if not query:
            return {'found': False, 'message': '查询词为空', 'results': []}

        query_words = self.preprocess_query(query)
        if not query_words:
            return {'found': False, 'message': '查询词无效', 'results': []}

        # 使用TF-IDF向量化
        try:
            vectorizer = TfidfVectorizer(max_features=50, min_df=1, max_df=0.8, ngram_range=(1, 2))

            # 准备文档列表（确保顺序与processed_files_content一致）
            doc_list = []
            file_names = []
            for file_name, content in processed_files_content.items():
                if not content.startswith("错误:"):
                    processed_content = self.preprocess_query(content)  # 使用相同的预处理
                    doc_list.append(' '.join(processed_content))
                    file_names.append(file_name)

            if not doc_list:
                return {'found': False, 'message': '没有有效文档进行TF-IDF分析', 'results': []}

            # 构建TF-IDF矩阵
            tfidf_matrix = vectorizer.fit_transform(doc_list)

            # 处理查询
            query_processed = ' '.join(query_words)
            query_vector = vectorizer.transform([query_processed])

            # 计算余弦相似度
            similarities = cosine_similarity(query_vector, tfidf_matrix).flatten()

            # 获取相似度排序的文档索引
            doc_indices = similarities.argsort()[::-1]

            results = []
            for idx in doc_indices:
                similarity_score = similarities[idx]
                if similarity_score > 0:  # 只返回有相似度的文档
                    file_name = file_names[idx]
                    content = processed_files_content[file_name]

                    # 找到匹配的查询词
                    matched_words = []
                    for word in query_words:
                        if word in content.lower():
                            matched_words.append(word)

                    preview = self.generate_content_preview(content, matched_words[0] if matched_words else query_words[0])

                    results.append({
                        'file_name': file_name,
                        'tfidf_score': round(similarity_score, 4),
                        'matched_words': matched_words,
                        'content_preview': preview,
                        'relevance': 'high' if similarity_score > 0.3 else 'medium' if similarity_score > 0.1 else 'low'
                    })

            return {
                'found': len(results) > 0,
                'query': query,
                'total_results': len(results),
                'results': results[:20],  # 限制返回前20个结果
                'message': f'TF-IDF检索找到 {len(results)} 个相关文档' if results else 'TF-IDF检索未找到相关文档'
            }

        except Exception as e:
            return {
                'found': False,
                'message': f'TF-IDF检索过程中发生错误: {str(e)}',
                'results': []
            }
    
    def lsi_search(self, query: str, inverted_index: Dict[str, Any], processed_files_content: Dict[str, str], file_word_count: Dict[str, Any], all_content: List[str], n_components: int = 100) -> Dict[str, Any]:
        """
        基于LSI（潜在语义索引）的检索
        Args:
            query: 搜索查询词
            inverted_index: 倒排索引
            processed_files_content: 文件内容字典
            file_word_count: 文件词频统计
            all_content: 所有文件内容列表
            n_components: LSI降维后的维度数
        Returns:
            检索结果字典
        """
        if not query:
            return {'found': False, 'message': '查询词为空', 'results': []}
    
        query_words = self.preprocess_query(query)
        if not query_words:
            return {'found': False, 'message': '查询词无效', 'results': []}
    
        try:
            # 使用TF-IDF向量化作为LSI的输入
            vectorizer = TfidfVectorizer(
                max_features=1000,  # 增加特征数量以获得更好的语义表示
                min_df=1, 
                max_df=0.8, 
                ngram_range=(1, 2),
                stop_words=None  # 我们已经有自己的停用词处理
            )
    
            # 准备文档列表
            doc_list = []
            file_names = []
            for file_name, content in processed_files_content.items():
                if not content.startswith("错误:"):
                    processed_content = self.preprocess_query(content)
                    doc_list.append(' '.join(processed_content))
                    file_names.append(file_name)
    
            if not doc_list:
                return {'found': False, 'message': '没有有效文档进行LSI分析', 'results': []}
    
            # 计算合适的LSI维度（不超过文档数量和词汇数量的最小值）
            n_components = min(n_components, len(doc_list) - 1, 1000)
            if n_components <= 0:
                n_components = min(50, len(doc_list) - 1)
    
            # 构建LSI pipeline
            lsi_pipeline = Pipeline([
                ('tfidf', vectorizer),
                ('svd', TruncatedSVD(n_components=n_components, random_state=42)),
                ('normalizer', Normalizer(copy=False))
            ])
    
            # 训练LSI模型并转换文档
            doc_lsi_matrix = lsi_pipeline.fit_transform(doc_list)
    
            # 处理查询并转换到LSI空间
            query_processed = ' '.join(query_words)
            query_lsi_vector = lsi_pipeline.transform([query_processed])
    
            # 计算LSI空间中的余弦相似度
            similarities = cosine_similarity(query_lsi_vector, doc_lsi_matrix).flatten()
    
            # 获取原始TF-IDF特征名称
            feature_names = vectorizer.get_feature_names_out()
            
            # 获取查询在原始TF-IDF空间的表示
            query_tfidf = vectorizer.transform([query_processed])
            query_tfidf_dense = query_tfidf.toarray()[0]
            
            # 找到查询中的关键词及其TF-IDF权重
            matched_features = [(feature_names[i], query_tfidf_dense[i]) 
                               for i in range(len(query_tfidf_dense)) 
                               if query_tfidf_dense[i] > 0]
    
            # 获取相似度排序的文档索引
            doc_indices = similarities.argsort()[::-1]
    
            results = []
            for idx in doc_indices:
                similarity_score = similarities[idx]
                if similarity_score > 0.01:  # 设置更低的阈值，因为LSI可能产生较小的相似度分数
                    file_name = file_names[idx]
                    content = processed_files_content[file_name]
    
                    # 获取该文档在LSI空间的表示
                    doc_lsi_vector = doc_lsi_matrix[idx]
                    
                    # 计算语义相关的词汇
                    doc_tfidf = vectorizer.transform([doc_list[idx]])
                    doc_tfidf_dense = doc_tfidf.toarray()[0]
                    
                    # 找到文档中与查询语义相关的词汇
                    semantic_matched_words = []
                    for word in query_words:
                        if word in feature_names:
                            word_idx = list(feature_names).index(word)
                            if doc_tfidf_dense[word_idx] > 0:
                                semantic_matched_words.append({
                                    'word': word,
                                    'tfidf_weight': round(doc_tfidf_dense[word_idx], 4),
                                    'frequency': processed_files_content[file_name].lower().count(word.lower())
                                })
    
                    # 获取LSI主题信息
                    lsi_components = self._get_top_lsi_terms(lsi_pipeline, feature_names, doc_lsi_vector)
    
                    preview = self.generate_content_preview(content, query_words[0])
    
                    results.append({
                        'file_name': file_name,
                        'lsi_similarity': round(similarity_score, 4),
                        'semantic_matched_words': semantic_matched_words,
                        'lsi_components': lsi_components[:5],  # 前5个主要LSI成分
                        'content_preview': preview,
                        'total_matched_terms': len(semantic_matched_words),
                        'algorithm': 'LSI',
                        'relevance': 'high' if similarity_score > 0.3 else 'medium' if similarity_score > 0.1 else 'low'
                    })
    
            return {
                'found': len(results) > 0,
                'query': query,
                'lsi_info': {
                    'n_components': n_components,
                    'query_features': matched_features[:10],  # 前10个查询特征
                    'total_query_terms': len(matched_features)
                },
                'total_results': len(results),
                'results': results[:20],  # 限制返回前20个结果
                'message': f'LSI检索找到 {len(results)} 个语义相关文档' if results else 'LSI检索未找到语义相关文档'
            }
    
        except Exception as e:
            return {
                'found': False,
                'message': f'LSI检索过程中发生错误: {str(e)}',
                'results': []
            }

        
class Visualization:
    def __init__(self):
        """初始化可视化类"""
        self.output_dir = 'server/src/pic/'
        self.ensure_output_dir()
        # 设置中文字体 - 添加字体存在性检查
        self._setup_matplotlib_fonts()
        # 停用词
        self.stopwords = self._load_stopwords()

    def _setup_matplotlib_fonts(self):
        """设置matplotlib中文字体"""
        import matplotlib.font_manager as fm
        font_path = 'server/src/font/simhei.ttf'  # 替换为您的实际字体文件路径
        try:
            if os.path.exists(font_path):
                # 添加字体
                fm.fontManager.addfont(font_path)
                 # 获取字体名称并设置
                font_prop = fm.FontProperties(fname=font_path)
                font_name = font_prop.get_name()

                plt.rcParams['font.sans-serif'] = [font_name, 'DejaVu Sans']
                plt.rcParams['axes.unicode_minus'] = False

                print(f"成功设置字体: {font_name}")
            else:
                print(f"字体文件不存在: {font_path}")
                plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
                plt.rcParams['axes.unicode_minus'] = False
            
        except Exception as e:
            print(f"字体设置错误: {e}")
            plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
            plt.rcParams['axes.unicode_minus'] = False
    def ensure_output_dir(self):
        """确保输出目录存在"""
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir, exist_ok=True)

    def _load_stopwords(self) -> List[str]:
        """加载停用词"""
        try:
            with open(stopwords_path, 'r', encoding='utf-8') as f:
                return [word.strip() for word in f.readlines()]
        except FileNotFoundError:
            return ['的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个']
    
    def preprocess_text_for_vis(self, text_data) -> List[str]:
        """为可视化预处理文本数据"""
        all_words = []

        if isinstance(text_data, dict):
            # 处理字典数据（文件内容字典）
            for content in text_data.values():
                if not content.startswith("错误:"):
                    words = self.preprocess_query(content)
                    all_words.extend(words)
        elif isinstance(text_data, list):
            # 处理列表数据
            for content in text_data:
                if isinstance(content, str) and not content.startswith("错误:"):
                    words = self.preprocess_query(content)
                    all_words.extend(words)
        elif isinstance(text_data, str):
            # 处理单个字符串
            words = self.preprocess_query(text_data)
            all_words.extend(words)

        return all_words

    def preprocess_query(self, query: str) -> List[str]:
        """预处理查询文本"""
        # 去除特殊字符，保留中文、英文、数字
        query = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s]', '', query)
        # 分词
        words = list(jieba.cut(query))
        # 去停用词和空词
        words = [word.strip() for word in words 
                if word.strip() and word.strip() not in self.stopwords and len(word.strip()) > 1]
        return words
    
    # 生成查询文章的词云图
    def generate_wordcloud(self, article_name: str, processed_files_content: Dict[str,str],filename: str = 'wordcloud', top_n: int = 50) -> Dict[str, Any]:
        """
        生成指定文章的词云图
    
        Args:
            article_name: 文章名（文件名）
            processed_files_content: 文件内容字典
            filename: 保存文件名
            top_n: 显示前N个词
    
        Returns:
            包含词频统计和保存路径的字典
        """
        try:
            if not article_name:
                return {
                    'success': False,
                    'message': '文章名不能为空',
                    'word_freq': {},
                    'save_path': None
                }
    
            # 查找指定的文章内容
            target_content = None
            matched_file_path = None
            
            # 遍历文件内容字典，查找匹配的文章
            for file_path, content in processed_files_content.items():
                if content.startswith("错误:"):
                    continue
                
                # 提取文件名进行匹配
                file_name = os.path.basename(file_path)
                if file_name == article_name or article_name in file_name:
                    target_content = content
                    matched_file_path = file_path
                    break
                
            if not target_content:
                return {
                    'success': False,
                    'message': f'未找到文章 "{article_name}"',
                    'word_freq': {},
                    'save_path': None
                }
    
            # 预处理文章文本内容
            all_words = self.preprocess_text_for_vis([target_content])
    
            if not all_words:
                return {
                    'success': False,
                    'message': f'文章 "{article_name}" 没有有效的文本数据生成词云',
                    'word_freq': {},
                    'save_path': None
                }
    
            # 统计词频
            word_freq = Counter(all_words)
            top_words = dict(word_freq.most_common(top_n))
    
            # 生成词云
            wordcloud = WordCloud(
                width=1200,
                height=800,
                background_color='white',
                max_words=top_n,
                colormap='viridis',
                relative_scaling=0.5,
                random_state=42,
                font_path='server/src/font/simhei.ttf'
            ).generate_from_frequencies(top_words)
    
            # 保存词云图
            # 使用文章名作为文件名的一部分，去除可能的特殊字符
            safe_article_name = re.sub(r'[^\w\-_\.]', '_', article_name)
            save_path = os.path.join(self.output_dir, 'cloud', f'{filename}_{safe_article_name}.png')
            
            # 确保保存目录存在
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
    
            plt.figure(figsize=(15, 10))
            plt.imshow(wordcloud, interpolation='bilinear')
            plt.axis('off')
            plt.title(f'文章 "{article_name}" 词云图 - 前{top_n}个高频词', fontsize=16, pad=20)
            plt.tight_layout()
            plt.savefig(save_path, dpi=300, bbox_inches='tight')
            plt.close()
    
            return {
                'success': True,
                'message': f'文章 "{article_name}" 的词云图已生成并保存到 {save_path}',
                'word_freq': top_words,
                'save_path': save_path,
                'total_words': len(all_words),
                'unique_words': len(word_freq),
                'article_name': article_name,
                'matched_file_path': matched_file_path
            }
    
        except Exception as e:
            return {
                'success': False,
                'message': f'生成文章 "{article_name}" 词云图时发生错误: {str(e)}',
                'word_freq': {},
                'save_path': None
            }

    # 生成查询文章的词向量图
    def generate_keyword_ranking(self, query: str, processed_files_content: Dict[str,str], file_word_count: Dict[str,int], filename: str = 'keyword_ranking', top_n: int = 10) -> Dict[str, Any]:
        """
        生成查询文章的关键词排行图

        Args:
            query: 查询词
            processed_files_content: 文件内容字典
            file_word_count: 文件词频统计
            filename: 保存文件名
            top_n: 显示前N个关键词

        Returns:
            包含关键词列表和保存路径的字典
        """
        try:
            # 预处理查询词
            query_words = self.preprocess_query(query)
            if not query_words:
                return {
                    'success': False,
                    'message': '查询词无效',
                    'keywords': [],
                    'save_path': None
                }

            # 收集包含查询词的文档的词频统计
            relevant_word_count = {}
            relevant_files = []

            for file_name, content in processed_files_content.items():
                if content.startswith("错误:"):
                    continue
                
                # 检查文档是否包含查询词
                content_words = self.preprocess_query(content)
                if any(word in content_words for word in query_words):
                    relevant_files.append(file_name)

                    # 合并该文档的词频统计
                    if file_name in file_word_count:
                        for word, count in file_word_count[file_name].items():
                            relevant_word_count[word] = relevant_word_count.get(word, 0) + count

            if not relevant_word_count:
                return {
                    'success': False,
                    'message': f'没有找到包含查询词 "{query}" 的文档词频数据',
                    'keywords': [],
                    'save_path': None
                }

            # 获取前N个关键词
            top_keywords = Counter(relevant_word_count).most_common(top_n)

            if not top_keywords:
                return {
                    'success': False,
                    'message': '没有足够的关键词生成排行',
                    'keywords': [],
                    'save_path': None
                }

            # 准备数据
            keywords = [item[0] for item in top_keywords]
            frequencies = [item[1] for item in top_keywords]

            # 生成排行图
            plt.figure(figsize=(12, 8))

            # 创建水平条形图
            bars = plt.barh(range(len(keywords)), frequencies, color='skyblue', alpha=0.8)

            # 设置标签
            plt.yticks(range(len(keywords)), keywords)
            plt.xlabel('词频', fontsize=12)
            plt.ylabel('关键词', fontsize=12)
            plt.title(f'查询词 "{query}" 相关文档关键词排行榜 - 前{top_n}名', fontsize=14, pad=20)

            # 在条形图上添加数值标签
            for i, (bar, freq) in enumerate(zip(bars, frequencies)):
                plt.text(bar.get_width() + max(frequencies) * 0.01, 
                        bar.get_y() + bar.get_height()/2, 
                        str(freq), 
                        ha='left', va='center', fontsize=10)

            # 反转y轴，使频率最高的在顶部
            plt.gca().invert_yaxis()

            # 调整布局
            plt.tight_layout()

            # 保存图片
            save_path = os.path.join(self.output_dir,'rank', f'{filename}_{query}.png')
            plt.savefig(save_path, dpi=300, bbox_inches='tight')
            plt.close()

            # 返回关键词列表
            keyword_list = [{'keyword': kw, 'frequency': freq, 'rank': i+1} 
                           for i, (kw, freq) in enumerate(top_keywords)]

            return {
                'success': True,
                'message': f'关键词排行图已生成并保存到 {save_path}',
                'keywords': keyword_list,
                'save_path': save_path,
                'relevant_files': relevant_files,
                'total_analyzed_files': len(relevant_files)
            }

        except Exception as e:
            return {
                'success': False,
                'message': f'生成关键词排行图时发生错误: {str(e)}',
                'keywords': [],
                'save_path': None
            }

    # 生成整个文件夹的文件的种类统计图
    def analyze_file_extensions(self, file_name: list, generate_chart: bool = True, filename: str = 'file_extensions') -> Dict[str, Any]:

        try:
            if not file_name:
                return {
                    'success': False,
                    'message': '文件名列表为空',
                    'extension_stats': [],
                    'save_path': None
                }

            # 提取文件后缀
            file_extensions = []
            for name in file_name:
                if isinstance(name, str):
                    ext = os.path.splitext(name)[1].lower()
                    if ext:
                        file_extensions.append(ext[1:])  # 去掉点号
                    else:
                        file_extensions.append('无后缀')

            if not file_extensions:
                return {
                    'success': False,
                    'message': '没有有效的文件数据进行分析',
                    'extension_stats': [],
                    'save_path': None
                }

            # 统计后缀频率
            ext_counter = Counter(file_extensions)
            extension_stats = [{'extension': ext, 'count': count, 'percentage': count/len(file_extensions)*100} 
                             for ext, count in ext_counter.most_common()]

            save_path = None

            if generate_chart and extension_stats:
                # 生成饼图
                plt.figure(figsize=(10, 8))

                extensions = [item['extension'] for item in extension_stats]
                counts = [item['count'] for item in extension_stats]

                # 设置颜色
                colors = plt.cm.Set3(np.linspace(0, 1, len(extensions)))

                # 创建饼图
                wedges, texts, autotexts = plt.pie(counts, labels=extensions, autopct='%1.1f%%', 
                                                  colors=colors, startangle=90)

                # 设置标题
                plt.title(f'文件类型分布统计 (总计: {len(file_extensions)} 个文件)', fontsize=14, pad=20)

                # 添加图例
                plt.legend(wedges, [f'{ext} ({count})' for ext, count in zip(extensions, counts)],
                          title="文件类型",
                          loc="center left",
                          bbox_to_anchor=(1, 0, 0.5, 1))

                plt.axis('equal')

                # 保存图片
                save_path = os.path.join(self.output_dir, 'extensions',f'{filename}.png')
                plt.savefig(save_path, dpi=300, bbox_inches='tight')
                plt.close()

            return {
                'success': True,
                'message': f'文件后缀分析完成' + (f'，图表已保存到 {save_path}' if save_path else ''),
                'extension_stats': extension_stats,
                'save_path': save_path,
                'total_files': len(file_extensions),
                'unique_extensions': len(ext_counter)
            }

        except Exception as e:
            return {
                'success': False,
                'message': f'分析文件后缀时发生错误: {str(e)}',
                'extension_stats': [],
                'save_path': None
            }

    # 统计全文件夹文件的词频前10
    def create_word_frequency_chart(self, word_dict: Dict[str, any], chart_type: str = 'bar', top_n: int = 20, filename: str = 'word_frequency') -> Dict[str, Any]:
        """
        统计全文件夹文件的词频前N名并生成图表

        Args:
            word_dict: 全文词频统计字典
            chart_type: 图表类型 ('bar', 'line', 'pie')
            top_n: 显示前N个词
            filename: 保存文件名

        Returns:
            图表生成结果
        """
        try:
            if not word_dict:
                return {
                    'success': False,
                    'message': '词频字典为空',
                    'save_path': None
                }

            # 获取前N个高频词
            if isinstance(word_dict, dict):
                # 如果已经是排序好的字典，直接取前N个
                top_words_items = list(word_dict.items())[:top_n]
            else:
                return {
                    'success': False,
                    'message': '词频数据格式不正确',
                    'save_path': None
                }

            if not top_words_items:
                return {
                    'success': False,
                    'message': '没有足够的词汇生成图表',
                    'save_path': None
                }

            words = [item[0] for item in top_words_items]
            frequencies = [item[1] for item in top_words_items]

            plt.figure(figsize=(12, 8))

            if chart_type == 'bar':
                bars = plt.bar(range(len(words)), frequencies, color='lightblue', alpha=0.8)
                plt.xticks(range(len(words)), words, rotation=45, ha='right')
                plt.ylabel('词频', fontsize=12)
                plt.xlabel('词汇', fontsize=12)

                # 在柱状图上添加数值标签
                for bar, freq in zip(bars, frequencies):
                    plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + max(frequencies) * 0.01,
                            str(freq), ha='center', va='bottom', fontsize=9)

            elif chart_type == 'line':
                plt.plot(range(len(words)), frequencies, marker='o', linewidth=2, markersize=6, color='blue')
                plt.xticks(range(len(words)), words, rotation=45, ha='right')
                plt.ylabel('词频', fontsize=12)
                plt.xlabel('词汇', fontsize=12)
                plt.grid(True, alpha=0.3)

                # 在数据点上添加数值标签
                for i, freq in enumerate(frequencies):
                    plt.text(i, freq + max(frequencies) * 0.02, str(freq), 
                            ha='center', va='bottom', fontsize=9)

            elif chart_type == 'pie':
                plt.pie(frequencies, labels=words, autopct='%1.1f%%', startangle=90)
                plt.axis('equal')

            plt.title(f'全文件夹词频统计 - {chart_type.upper()}图 (前{top_n}个)', fontsize=14, pad=20)
            plt.tight_layout()

            # 保存图片
            save_path = os.path.join(self.output_dir,'frequency', f'{filename}_{chart_type}.png')
            plt.savefig(save_path, dpi=300, bbox_inches='tight')
            plt.close()

            return {
                'success': True,
                'message': f'{chart_type.upper()}图已生成并保存到 {save_path}',
                'save_path': save_path,
                'chart_type': chart_type,
                'top_words': top_words_items,
                'total_unique_words': len(word_dict),
                'analyzed_words': len(top_words_items)
            }

        except Exception as e:
            return {
                'success': False,
                'message': f'生成{chart_type}图时发生错误: {str(e)}',
                'save_path': None
            }


                                       

        
